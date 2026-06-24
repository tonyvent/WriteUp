"""Turn a recorded session into a readable step-by-step write-up.

Produces Markdown and a standalone HTML report by default; can also emit a
Word document if python-docx is installed. Steps are grouped under headings
whenever the active application changes, so a process that spans (say)
AutoCAD and a browser reads naturally.
"""
from __future__ import annotations

import base64
import html
from pathlib import Path
from typing import Iterable

from .models import Event, Session


# --------------------------------------------------------------------- phrasing
def describe(event: Event) -> str:
    """A human-readable sentence for a single event."""
    app = event.app or "the application"
    if event.kind == "click":
        btn = event.button or "left"
        verb = "Click" if btn == "left" else f"{btn.capitalize()}-click"
        return f"{verb} at the highlighted location in {app}."
    if event.kind == "type":
        return f"Type `{event.text}`."
    if event.kind == "key":
        pretty = {"enter": "Enter", "tab": "Tab", "esc": "Esc"}.get(
            event.text, event.text
        )
        return f"Press **{pretty}**."
    if event.kind == "scroll":
        return f"Scroll in {app}."
    if event.kind == "note":
        return event.text or "Note."
    return event.kind


def _step_events(session: Session) -> Iterable[Event]:
    """Everything except bare window markers (those become headings)."""
    return [e for e in session.events if e.kind != "window"]


def _title_for(session: Session) -> str:
    if session.meta.title:
        return session.meta.title
    return session.name.replace("-", " ").replace("_", " ").strip().title()


def _doc_date(session: Session) -> str:
    """Document date: explicit meta date, else the recorded date (date only)."""
    if session.meta.date:
        return session.meta.date
    if session.started:
        return session.started.split("T", 1)[0]
    return ""


def _meta_bits(session: Session) -> list[tuple[str, str]]:
    """Ordered (label, value) pairs to show in the title block."""
    m = session.meta
    bits: list[tuple[str, str]] = []
    if m.author:
        bits.append(("Author", m.author))
    date = _doc_date(session)
    if date:
        bits.append(("Date", date))
    if m.revision:
        bits.append(("Revision", m.revision))
    if m.department:
        bits.append(("Department", m.department))
    return bits


def _prepare_logo(session: Session) -> str | None:
    """Copy the logo into the session's assets/ folder; return the rel path."""
    if not session.meta.logo:
        return None
    src = Path(session.meta.logo).expanduser()
    if not src.exists():
        return None
    assets = session.folder / "assets"
    assets.mkdir(parents=True, exist_ok=True)
    dest = assets / f"logo{src.suffix.lower()}"
    if src.resolve() != dest.resolve():
        dest.write_bytes(src.read_bytes())
    return f"assets/{dest.name}"


# -------------------------------------------------------------------- markdown
def to_markdown(session: Session) -> str:
    m = session.meta
    lines: list[str] = []

    logo = _prepare_logo(session)
    if logo:
        lines += [f"![logo]({logo})", ""]
    if m.company:
        header = f"**{m.company}**"
        if m.department:
            header += f" — {m.department}"
        lines += [header, ""]

    lines += [f"# {_title_for(session)}", ""]
    if m.subtitle:
        lines += [f"_{m.subtitle}_", ""]

    bits = _meta_bits(session)
    if bits:
        lines.append(" | ".join(f"**{k}:** {v}" for k, v in bits))
        lines.append("")
    lines.append("---")
    lines.append("")

    step_no = 0
    current_app = None
    for event in session.events:
        if event.kind == "window":
            if event.app and event.app != current_app:
                current_app = event.app
                lines.append(f"\n## In {event.app}\n")
            continue
        step_no += 1
        lines.append(f"**{step_no}.** {describe(event)}")
        if event.screenshot:
            lines.append("")
            lines.append(f"![Step {step_no}]({event.screenshot})")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


# ------------------------------------------------------------------------ html
_CSS = """
:root { --accent:#ff4c1f; --ink:#1a1a1a; --muted:#6b6b6b; --line:#e7e7e7; }
* { box-sizing:border-box; }
body { margin:0; font:16px/1.6 -apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;
       color:var(--ink); background:#fafafa; }
.wrap { max-width:820px; margin:0 auto; padding:48px 24px 96px; }
.brandbar { display:flex; align-items:center; justify-content:space-between; gap:16px;
            padding-bottom:16px; border-bottom:2px solid var(--accent); margin-bottom:28px; }
.brandbar .co { font-weight:700; font-size:18px; letter-spacing:-.01em; }
.brandbar .co small { display:block; font-weight:500; font-size:13px; color:var(--muted); }
.brandbar img.logo { height:44px; width:auto; border:0; border-radius:0; margin:0; }
h1 { font-size:32px; margin:0 0 4px; letter-spacing:-.02em; }
.subtitle { color:var(--muted); font-size:17px; margin:0 0 16px; }
.metarow { display:flex; flex-wrap:wrap; gap:8px 24px; margin:0 0 28px;
           padding:12px 0; border-top:1px solid var(--line); border-bottom:1px solid var(--line); }
.metarow .item { font-size:14px; }
.metarow .item span { color:var(--muted); text-transform:uppercase; letter-spacing:.06em;
                      font-size:11px; display:block; }
h2 { font-size:14px; text-transform:uppercase; letter-spacing:.08em; color:var(--accent);
     margin:40px 0 8px; }
.step { display:flex; gap:16px; padding:16px 0; border-top:1px solid var(--line); }
.num { flex:0 0 32px; height:32px; border-radius:50%; background:var(--accent); color:#fff;
       font-weight:700; display:flex; align-items:center; justify-content:center; font-size:14px; }
.body { flex:1; min-width:0; }
.body p { margin:4px 0 0; }
.body code { background:#f0f0f0; padding:1px 6px; border-radius:4px; font-size:.92em; }
img { width:100%; border:1px solid var(--line); border-radius:8px; margin-top:12px; display:block; }
footer { margin-top:48px; color:var(--muted); font-size:13px; text-align:center; }
@media print {
  body { background:#fff; }
  .wrap { max-width:none; padding:0 12px; }
  .step, .metarow, .brandbar { break-inside:avoid; }
  img { break-inside:avoid; }
  footer { display:none; }
}
"""


def to_html(session: Session, embed_images: bool = False) -> str:
    m = session.meta
    parts: list[str] = [
        "<!doctype html><html lang='en'><head><meta charset='utf-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1'>",
        f"<title>{html.escape(_title_for(session))}</title>",
        f"<style>{_CSS}</style></head><body><div class='wrap'>",
    ]

    # Brand bar: company / department on the left, logo on the right.
    logo_rel = _prepare_logo(session)
    if m.company or logo_rel:
        parts.append("<div class='brandbar'>")
        co = ""
        if m.company:
            co = f"<div class='co'>{html.escape(m.company)}"
            if m.department:
                co += f"<small>{html.escape(m.department)}</small>"
            co += "</div>"
        parts.append(co or "<div></div>")
        if logo_rel:
            src = _image_src(session, logo_rel, embed_images)
            if src:
                parts.append(f"<img class='logo' src='{src}' alt='logo'>")
        parts.append("</div>")

    parts.append(f"<h1>{html.escape(_title_for(session))}</h1>")
    if m.subtitle:
        parts.append(f"<p class='subtitle'>{html.escape(m.subtitle)}</p>")

    bits = _meta_bits(session)
    if bits:
        parts.append("<div class='metarow'>")
        for label, value in bits:
            parts.append(
                f"<div class='item'><span>{html.escape(label)}</span>"
                f"{html.escape(value)}</div>"
            )
        parts.append("</div>")

    step_no = 0
    current_app = None
    for event in session.events:
        if event.kind == "window":
            if event.app and event.app != current_app:
                current_app = event.app
                parts.append(f"<h2>In {html.escape(event.app)}</h2>")
            continue
        step_no += 1
        desc = _md_inline_to_html(describe(event))
        parts.append("<div class='step'>")
        parts.append(f"<div class='num'>{step_no}</div><div class='body'>")
        parts.append(f"<p>{desc}</p>")
        if event.screenshot:
            src = _image_src(session, event.screenshot, embed_images)
            if src:
                parts.append(f"<img src='{src}' alt='Step {step_no}'>")
        parts.append("</div></div>")

    parts.append("<footer>Generated by ProcessScribe</footer>")
    parts.append("</div></body></html>")
    return "".join(parts)


def _md_inline_to_html(text: str) -> str:
    text = html.escape(text)
    # Tiny inline markdown: **bold** and `code`.
    import re

    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    return text


def _image_src(session: Session, rel: str, embed: bool) -> str | None:
    if not embed:
        return rel
    path = session.folder / rel
    if not path.exists():
        return None
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{data}"


# ------------------------------------------------------------------------ docx
def to_docx(session: Session, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    m = session.meta
    doc = Document()

    logo_rel = _prepare_logo(session)
    if logo_rel:
        p = doc.add_paragraph()
        p.add_run().add_picture(str(session.folder / logo_rel), height=Inches(0.5))
    if m.company:
        run = doc.add_paragraph().add_run(m.company)
        run.bold = True
        run.font.size = Pt(14)
        if m.department:
            sub = doc.add_paragraph().add_run(m.department)
            sub.italic = True
            sub.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

    doc.add_heading(_title_for(session), level=0)
    if m.subtitle:
        doc.add_paragraph(m.subtitle).runs[0].italic = True

    bits = _meta_bits(session)
    if bits:
        meta_p = doc.add_paragraph()
        meta_p.add_run("   ".join(f"{k}: {v}" for k, v in bits)).font.size = Pt(10)

    step_no = 0
    current_app = None
    for event in session.events:
        if event.kind == "window":
            if event.app and event.app != current_app:
                current_app = event.app
                doc.add_heading(f"In {event.app}", level=1)
            continue
        step_no += 1
        doc.add_paragraph(
            f"{step_no}. {describe(event).replace('`', '').replace('**', '')}"
        )
        if event.screenshot:
            img = session.folder / event.screenshot
            if img.exists():
                doc.add_picture(str(img), width=Inches(6))
    doc.save(str(out_path))
    return out_path


# ------------------------------------------------------------------------- pdf
_ACCENT = (255, 76, 31)
_MUTED = (107, 107, 107)


def _latin1(text: str) -> str:
    """fpdf core fonts are latin-1 only; map common unicode and drop the rest."""
    repl = {
        "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u00a0": " ",
        "\u2022": "-", "\u2192": "->",
    }
    for bad, good in repl.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "replace").decode("latin-1")


def _plain(event: Event) -> str:
    return describe(event).replace("`", "").replace("**", "")


def to_pdf(session: Session, out_path: Path, page_size: str = "Letter") -> Path:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    from PIL import Image

    m = session.meta
    logo_rel = _prepare_logo(session)

    class Report(FPDF):
        def footer(self) -> None:
            self.set_y(-36)
            self.set_font("Helvetica", size=8)
            self.set_text_color(*_MUTED)
            left = _latin1(m.company or "Generated by ProcessScribe")
            self.cell(0, 10, left)
            self.cell(0, 10, f"Page {self.page_no()}", align="R")

    pdf = Report(unit="pt", format=page_size)
    pdf.set_margins(54, 54, 54)
    pdf.set_auto_page_break(True, margin=54)
    pdf.add_page()
    content_w = pdf.w - pdf.l_margin - pdf.r_margin

    # --- header / branding ------------------------------------------------
    if logo_rel:
        lp = session.folder / logo_rel
        lw, lh = Image.open(lp).size
        disp_h = 34
        disp_w = disp_h * lw / lh
        pdf.image(str(lp), x=pdf.w - pdf.r_margin - disp_w, y=pdf.t_margin, h=disp_h)
    if m.company:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(26, 26, 26)
        pdf.cell(content_w, 16, _latin1(m.company), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if m.department:
            pdf.set_font("Helvetica", size=10)
            pdf.set_text_color(*_MUTED)
            pdf.cell(content_w, 13, _latin1(m.department),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(8)
    pdf.set_draw_color(*_ACCENT)
    pdf.set_line_width(2)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(14)

    # --- title block ------------------------------------------------------
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(26, 26, 26)
    pdf.multi_cell(content_w, 28, _latin1(_title_for(session)),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if m.subtitle:
        pdf.set_font("Helvetica", "I", 12)
        pdf.set_text_color(*_MUTED)
        pdf.multi_cell(content_w, 16, _latin1(m.subtitle),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    bits = _meta_bits(session)
    if bits:
        pdf.ln(4)
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(60, 60, 60)
        line = "    ".join(f"{k}: {v}" for k, v in bits)
        pdf.multi_cell(content_w, 14, _latin1(line),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(12)

    # --- steps ------------------------------------------------------------
    step_no = 0
    current_app = None
    for event in session.events:
        if event.kind == "window":
            if event.app and event.app != current_app:
                current_app = event.app
                pdf.ln(6)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(*_ACCENT)
                pdf.multi_cell(content_w, 16, _latin1(f"IN {event.app.upper()}"),
                               new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)
            continue
        step_no += 1
        pdf.set_font("Helvetica", size=11)
        pdf.set_text_color(26, 26, 26)
        pdf.multi_cell(content_w, 16, _latin1(f"{step_no}.  {_plain(event)}"),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if event.screenshot:
            img = session.folder / event.screenshot
            if img.exists():
                iw, ih = Image.open(img).size
                draw_w = min(content_w, 460)
                draw_h = draw_w * ih / iw
                if pdf.get_y() + draw_h > pdf.h - pdf.b_margin:
                    pdf.add_page()
                y0 = pdf.get_y() + 4
                pdf.image(str(img), x=pdf.l_margin, y=y0, w=draw_w, h=draw_h)
                pdf.set_y(y0 + draw_h + 10)
        else:
            pdf.ln(4)

    pdf.output(str(out_path))
    return out_path



# ----------------------------------------------------------------------- write
_DEP_HINT = {"docx": "python-docx", "pdf": "fpdf2"}


def _resolve_targets(fmt: str) -> list[str]:
    fmt = (fmt or "all").lower()
    if fmt == "all":
        return ["md", "html", "docx", "pdf"]
    aliases = {"markdown": "md", "word": "docx"}
    return [aliases.get(fmt, fmt)]


def generate(
    session: Session,
    fmt: str = "all",
    embed: bool = False,
    page_size: str = "Letter",
) -> list[Path]:
    """Write the requested format(s) into the session folder.

    For an explicit single format, a missing optional dependency raises
    ImportError (so the caller can show an install hint). For "all", any
    format whose dependency is missing is skipped silently.
    """
    written: list[Path] = []
    folder = session.folder
    explicit = (fmt or "all").lower() != "all"

    for target in _resolve_targets(fmt):
        try:
            if target == "md":
                p = folder / "report.md"
                p.write_text(to_markdown(session), encoding="utf-8")
            elif target == "html":
                p = folder / "report.html"
                p.write_text(to_html(session, embed_images=embed), encoding="utf-8")
            elif target == "docx":
                p = to_docx(session, folder / "report.docx")
            elif target == "pdf":
                p = to_pdf(session, folder / "report.pdf", page_size=page_size)
            else:
                continue
            written.append(p)
        except ImportError:
            if explicit:
                raise
            continue
    return written
